from telegram import Update
from telegram.ext import ContextTypes, ConversationHandler
from db.cotizaciones_db import obtener_datos_cliente, insertar_cotizacion, buscar_cotizacion_id
from datetime import datetime
from db.conexion import obtener_conexion
import os
from docx import Document
from docx2pdf import convert




IDENTIFICACION, DATOS_COTIZACION = 20, 21

async def iniciar_cotizacion(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text("🧾 Por favor envíame la identificación del cliente.")
    return IDENTIFICACION

async def recibir_identificacion(update: Update, context: ContextTypes.DEFAULT_TYPE):
    identificacion = update.message.text.strip()
    cliente = obtener_datos_cliente(identificacion)

    if cliente:
        nombres, apellidos, correo = cliente
        context.user_data["cliente"] = {
            "nombre": f"{nombres} {apellidos}",
            "identificacion": identificacion,
            "correo": correo,
        }

        await update.message.reply_text(
            f"✅ Cliente encontrado:\n"
            f"Nombre: {nombres} {apellidos}\n"
            f"Correo: {correo}\n\n"
            f"🧾 Ahora envíame los datos restantes de la cotización en formato:"
        )

        ejemplo = (
            "fecha: 2025-06-21\n"
            "marca: Toyota\n"
            "referencia: Hilux\n"
            "asesor_chasis: Juan Pérez\n"
            "asesor_inducar: María Gómez\n"
            "carroceria: Pick-up\n"
            "tipo_carroceria: Doble cabina\n"
            "medidas_internas: 2.5x1.5x1.2\n"
            "medidas_externas: 5.3x2.0x1.8\n"
            "descrip_medidas: Medidas aproximadas de la carrocería\n"
            "forro_interno: Sí\n"
            "forro_externo: No\n"
            "vigas: 4\n"
            "puente: 2\n"
            "piso: Antideslizante\n"
            "marco: Reforzado\n"
            "parales: 6\n"
            "compuertas: 2\n"
            "puertas: 4\n"
            "color: Blanco\n"
            "guarda_polvos: Sí\n"
            "adicional: Alarma\n"
            "descrip_adicional: Sistema de alarma con sensor de movimiento\n"
            "anclaje: Sí\n"
            "bomper: Reforzado delantero y trasero\n"
            "aseguramiento: Seguro antivuelco\n"
            "mecanismo_cierre: Manual\n"
            "impermeabilizacion: Sí\n"
            "accesorios_acero: Barras laterales\n"
            "accesorios: Extintor\n"
            "observaciones: Entrega en 15 días hábiles\n"
            "forma_pago: Crédito\n"
            "valor_total: 45000\n"
            "subtotal: 37815\n"
            "iva_19: 7185\n"
            "fecha_cierre: 2025-07-01\n"
            "largo_externo: 5.3\n"
            "ancho_externo: 2.0\n"
            "alto_externo: 1.8\n"
            "largo_internas: 2.5\n"
            "ancho_internas: 1.5\n"
            "alto_internas: 1.2"
        )

        await update.message.reply_text(ejemplo)
        return DATOS_COTIZACION

    else:
        await update.message.reply_text("❌ No se encontró el cliente. Intenta de nuevo.")
        return IDENTIFICACION

def obtener_consecutivo(año, mes, carroceria, tipo_carroceria) -> int:
    try:
        conn = obtener_conexion()
        cur = conn.cursor()
        
        prefijo = f"{carroceria[0].upper()}{tipo_carroceria[0].upper()}{año}{mes:02d}"

        cur.execute("""
            SELECT MAX(CAST(SUBSTRING(no_cotizacion, 9) AS INTEGER))
            FROM ventas.cotizaciones
            WHERE SUBSTRING(no_cotizacion, 1, 8) = %s
        """, (prefijo,))
        
        max_consecutivo = cur.fetchone()[0]
        cur.close()
        conn.close()

        if max_consecutivo is None:
            return 101
        else:
            return max_consecutivo + 1

    except Exception as e:
        print("Error obteniendo consecutivo:", e)
        return 101  # valor de fallback
    
async def recibir_datos_cotizacion(update: Update, context: ContextTypes.DEFAULT_TYPE):
    texto = update.message.text
    datos = {}
    for linea in texto.split("\n"):
        if ":" in linea:
            k, v = linea.split(":", 1)
            datos[k.strip().lower()] = v.strip()

    cliente = context.user_data["cliente"]

    # Parsear año y mes de la fecha
    fecha_str = datos.get('fecha')
    if fecha_str:
        fecha_dt = datetime.strptime(fecha_str, "%Y-%m-%d")
        año = fecha_dt.year
        mes = fecha_dt.month
    else:
        fecha_dt = datetime.now()
        año = fecha_dt.year
        mes = fecha_dt.month

    carroceria = datos.get('carroceria', '').upper()
    tipo_carroceria = datos.get('tipo_carroceria', '').upper()
    letra_c = carroceria[0] if carroceria else 'X'
    letra_t = tipo_carroceria[0] if tipo_carroceria else 'X'

    consecutivo = obtener_consecutivo(año, mes, carroceria, tipo_carroceria)
    codigo_cotizacion = f"{letra_c}{letra_t}{año}{mes:02d}{consecutivo}"


    datos_db = {
        'no_cotizacion': codigo_cotizacion,
        'nombre': cliente['nombre'],
        'identificacion': cliente['identificacion'],
        'correo': cliente['correo'],
        'fecha': datos.get('fecha'),
        'marca': datos.get('marca'),
        'referencia': datos.get('referencia'),
        'asesor_chasis': datos.get('asesor_chasis'),
        'asesor_inducar': datos.get('asesor_inducar'),
        'carroceria': datos.get('carroceria'),
        'tipo_carroceria': datos.get('tipo_carroceria'),
        'medidas_internas': datos.get('medidas_internas'),
        'medidas_externas': datos.get('medidas_externas'),
        'descrip_medidas': datos.get('descrip_medidas'),
        'forro_interno': datos.get('forro_interno'),
        'forro_externo': datos.get('forro_externo'),
        'vigas': datos.get('vigas'),
        'puente': datos.get('puente'),
        'piso': datos.get('piso'),
        'marco': datos.get('marco'),
        'parales': datos.get('parales'),
        'compuertas': datos.get('compuertas'),
        'puertas': datos.get('puertas'),
        'color': datos.get('color'),
        'guarda_polvos': datos.get('guarda_polvos'),
        'adicional': datos.get('adicional'),
        'descrip_adicional': datos.get('descrip_adicional'),
        'anclaje': datos.get('anclaje'),
        'bomper': datos.get('bomper'),
        'aseguramiento': datos.get('aseguramiento'),
        'mecanismo_cierre': datos.get('mecanismo_cierre'),
        'impermeabilizacion': datos.get('impermeabilizacion'),
        'accesorios_acero': datos.get('accesorios_acero'),
        'accesorios': datos.get('accesorios'),
        'observaciones': datos.get('observaciones'),
        'forma_pago': datos.get('forma_pago'),
        'valor_total': float(datos.get('valor_total') or 0),
        'subtotal': float(datos.get('subtotal') or 0),
        'iva_19': float(datos.get('iva_19') or 0),
        'fecha_cierre': datos.get('fecha_cierre'),
        'largo_externo': float(datos.get('largo_externo') or 0),
        'ancho_externo': float(datos.get('ancho_externo') or 0),
        'alto_externo': float(datos.get('alto_externo') or 0),
        'largo_internas': float(datos.get('largo_internas') or 0),
        'ancho_internas': float(datos.get('ancho_internas') or 0),
        'alto_internas': float(datos.get('alto_internas') or 0),
    }

    exito = insertar_cotizacion(datos_db)
    if exito:
        await update.message.reply_text(f"✅ Cotización guardada exitosamente.\nNúmero de cotización: {codigo_cotizacion}")
    else:
        await update.message.reply_text("❌ Hubo un error al guardar la cotización.")

    
    return ConversationHandler.END

#ENVIAR COTIZACION

CARPETA_FORMATOS = "/Users/andresgalvis/Library/CloudStorage/OneDrive-SharedLibraries-industriasycarrocerias/INDUCAR SAS - Documentos/Ventas/Informacion_Ventas"
CARPETA_BASE_PDF = "/Users/andresgalvis/Library/CloudStorage/OneDrive-SharedLibraries-industriasycarrocerias/INDUCAR SAS - Documentos/Ventas/Cotizaciones"

def reemplazar_marcadores(doc: Document, datos: dict):
    # Reemplazo en párrafos normales
    for p in doc.paragraphs:
        for k, v in datos.items():
            marcador = f"{{{{{k}}}}}"
            if marcador in p.text:
                p.text = p.text.replace(marcador, str(v))

    # Reemplazo en tablas
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for p in celda.paragraphs:
                    for k, v in datos.items():
                        marcador = f"{{{{{k}}}}}"
                        if marcador in p.text:
                            p.text = p.text.replace(marcador, str(v))

async def cotizacion_pdf(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text("📄 Envíame el número de cotización (ej: PT202507123) para generar el PDF.")
    return 23  # Nuevo estado para que después se procese el número



async def generar_pdf(update: Update, context: ContextTypes.DEFAULT_TYPE):
    no_cotizacion = update.message.text.strip()
    print(f"[DEBUG] Recibido número de cotización: {no_cotizacion}")

    datos = buscar_cotizacion_id(no_cotizacion)
    print(f"[DEBUG] Datos encontrados: {datos}")

    if not datos:
        await update.message.reply_text("❌ No se encontró la cotización.")
        return

    try:
        carroceria = datos.get("carroceria", "Generico").strip().capitalize()
        tipo_carroceria = datos.get("tipo_carroceria", "Generico").strip().capitalize()
        fecha = datos["fecha"]
        año = fecha.year
        mes = fecha.month

        nombre_plantilla = f"{carroceria} {tipo_carroceria}.docx"
        ruta_plantilla = os.path.join(CARPETA_FORMATOS, "Cotizaciones", nombre_plantilla)
        print(f"[DEBUG] Ruta plantilla: {ruta_plantilla}")

        if not os.path.exists(ruta_plantilla):
            await update.message.reply_text(f"❌ No se encontró la plantilla: {nombre_plantilla}")
            return

        doc = Document(ruta_plantilla)
        print(f"[DEBUG] Plantilla cargada correctamente.")

        
        reemplazar_marcadores(doc, datos)


        carpeta_destino = os.path.join(CARPETA_BASE_PDF, str(año), f"{mes:02d}")
        os.makedirs(carpeta_destino, exist_ok=True)
        ruta_word = os.path.join(carpeta_destino, f"{no_cotizacion}.docx")
        doc.save(ruta_word)
        print(f"[DEBUG] Documento Word guardado en: {ruta_word}")

        ruta_pdf = os.path.join(carpeta_destino, f"{no_cotizacion}.pdf")
        convert(ruta_word, ruta_pdf)
        print(f"[DEBUG] PDF generado en: {ruta_pdf}")

        with open(ruta_pdf, "rb") as archivo:
            await update.message.reply_document(document=archivo, filename=f"{no_cotizacion}.pdf")

        await update.message.reply_text("✅ Cotización generada y enviada con éxito.")

    except Exception as e:
        await update.message.reply_text(f"❌ Error generando PDF: {e}")
        print(f"[ERROR] {e}")

    return ConversationHandler.END


async def buscar_cotizacion(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text("🔍 Por favor envíame la NO_COTIZACION de la cotizacion a buscar:")
    return 22  # NUEVO ESTADO

async def mostrar_cotizacion(update: Update, context: ContextTypes.DEFAULT_TYPE):
    no_cotizacion = update.message.text.strip()
    cotizaciones = buscar_cotizacion_id(no_cotizacion)

    if cotizaciones:
        texto = "\n".join(f"{k}: {v}" for k, v in cotizaciones.items())
        await update.message.reply_text(f"✅ Cotizacion encontrado:\n\n{texto}")
    else:
        await update.message.reply_text("❌ No se encontró ningún cotizacion con esa NO. Cotizacion.")

    return ConversationHandler.END
